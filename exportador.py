# exportador.py
# Exporta contratos, presupuestos e informes a PDF y DOCX profesionales

from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ─── Paleta de colores ─────────────────────────────────────────────────────────
AZUL_PRINCIPAL = colors.HexColor("#1E3A5F")
AZUL_CLARO = colors.HexColor("#2E6DA4")
GRIS_LINEA = colors.HexColor("#D0D7E3")
GRIS_FONDO = colors.HexColor("#F4F6FA")
NEGRO = colors.HexColor("#1A1A2E")
BLANCO = colors.white


# ─── Utilidades comunes ────────────────────────────────────────────────────────

def _parsear_texto(texto: str) -> list[dict]:
    """Convierte texto plano con marcadores comunes en bloques tipados."""
    bloques = []
    lineas = texto.splitlines()
    for linea in lineas:
        linea_strip = linea.strip()
        if not linea_strip:
            bloques.append({"tipo": "espacio"})
            continue
        # Detectar encabezados tipo "CLÁUSULA", "ARTÍCULO", sección en mayúsculas
        if re.match(r'^(CLÁUSULA|ARTÍCULO|SECCIÓN|ANEXO|PARTE|CAPÍTULO)\s', linea_strip, re.IGNORECASE):
            bloques.append({"tipo": "h2", "texto": linea_strip})
        elif linea_strip.isupper() and len(linea_strip) > 4 and len(linea_strip) < 80:
            bloques.append({"tipo": "h1", "texto": linea_strip})
        elif linea_strip.startswith(("- ", "• ", "* ")):
            bloques.append({"tipo": "bullet", "texto": linea_strip[2:]})
        elif re.match(r'^\d+[\.\)]\s', linea_strip):
            bloques.append({"tipo": "numerado", "texto": linea_strip})
        else:
            bloques.append({"tipo": "parrafo", "texto": linea_strip})
    return bloques


# ─── EXPORTAR A PDF ─────────────────────────────────────────────────────────────

def exportar_pdf(contenido: str, tipo_doc: str = "documento", metadata: dict = None) -> bytes:
    """
    Genera un PDF profesional a partir del texto generado por IA.

    Args:
        contenido: Texto del documento generado
        tipo_doc: "contrato" | "presupuesto" | "informe"
        metadata: dict con claves opcionales: empresa, cliente, fecha, numero

    Returns:
        bytes del PDF listo para descargar
    """
    buffer = BytesIO()
    metadata = metadata or {}

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    estilos = getSampleStyleSheet()

    # Estilos personalizados
    estilo_h1 = ParagraphStyle(
        "H1Custom",
        parent=estilos["Normal"],
        fontSize=13,
        fontName="Helvetica-Bold",
        textColor=AZUL_PRINCIPAL,
        spaceAfter=6,
        spaceBefore=14,
        leading=16,
    )
    estilo_h2 = ParagraphStyle(
        "H2Custom",
        parent=estilos["Normal"],
        fontSize=11,
        fontName="Helvetica-Bold",
        textColor=AZUL_CLARO,
        spaceAfter=4,
        spaceBefore=10,
        leading=14,
    )
    estilo_parrafo = ParagraphStyle(
        "ParrafoCustom",
        parent=estilos["Normal"],
        fontSize=10,
        fontName="Helvetica",
        textColor=NEGRO,
        spaceAfter=4,
        spaceBefore=2,
        leading=15,
        alignment=TA_JUSTIFY,
    )
    estilo_bullet = ParagraphStyle(
        "BulletCustom",
        parent=estilos["Normal"],
        fontSize=10,
        fontName="Helvetica",
        textColor=NEGRO,
        leftIndent=20,
        spaceAfter=3,
        leading=14,
        bulletIndent=8,
    )

    historia = []

    # ── Cabecera del documento ──────────────────────────────────────────────
    # Barra superior de color
    historia.append(
        Table(
            [[Paragraph(
                f'<font color="white"><b>{tipo_doc.upper()}</b></font>',
                ParagraphStyle("cab", fontSize=14, fontName="Helvetica-Bold",
                               textColor=BLANCO, alignment=TA_CENTER)
            )]],
            colWidths=[doc.width],
            style=TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), AZUL_PRINCIPAL),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
            ])
        )
    )
    historia.append(Spacer(1, 0.4 * cm))

    # Metadatos en tabla de dos columnas si existen
    filas_meta = []
    if metadata.get("empresa"):
        filas_meta.append(["Empresa:", metadata["empresa"]])
    if metadata.get("cliente"):
        filas_meta.append(["Cliente:", metadata["cliente"]])
    if metadata.get("fecha"):
        filas_meta.append(["Fecha:", metadata["fecha"]])
    if metadata.get("numero"):
        filas_meta.append(["Número:", metadata["numero"]])

    if filas_meta:
        tabla_meta = Table(
            filas_meta,
            colWidths=[3.5 * cm, doc.width - 3.5 * cm],
            style=TableStyle([
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("TEXTCOLOR", (0, 0), (0, -1), AZUL_CLARO),
                ("TEXTCOLOR", (1, 0), (1, -1), NEGRO),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("BACKGROUND", (0, 0), (-1, -1), GRIS_FONDO),
                ("BOX", (0, 0), (-1, -1), 0.5, GRIS_LINEA),
            ])
        )
        historia.append(tabla_meta)
        historia.append(Spacer(1, 0.5 * cm))

    historia.append(HRFlowable(width="100%", thickness=1.5, color=AZUL_PRINCIPAL))
    historia.append(Spacer(1, 0.4 * cm))

    # ── Cuerpo del documento ────────────────────────────────────────────────
    bloques = _parsear_texto(contenido)

    for bloque in bloques:
        t = bloque["tipo"]
        if t == "espacio":
            historia.append(Spacer(1, 0.25 * cm))
        elif t == "h1":
            historia.append(Paragraph(bloque["texto"], estilo_h1))
            historia.append(HRFlowable(width="100%", thickness=0.5, color=GRIS_LINEA))
        elif t == "h2":
            historia.append(Paragraph(bloque["texto"], estilo_h2))
        elif t == "bullet":
            historia.append(Paragraph(f"&#8226;  {bloque['texto']}", estilo_bullet))
        elif t == "numerado":
            historia.append(Paragraph(bloque["texto"], estilo_bullet))
        else:
            historia.append(Paragraph(bloque["texto"], estilo_parrafo))

    # ── Pie de firma ────────────────────────────────────────────────────────
    historia.append(Spacer(1, 1.5 * cm))
    historia.append(HRFlowable(width="100%", thickness=0.5, color=GRIS_LINEA))
    historia.append(Spacer(1, 0.3 * cm))

    ancho_col = doc.width / 2 - 0.5 * cm
    firma_data = [
        [
            Paragraph("Firma y sello<br/><br/><br/>___________________________<br/>Representante empresa",
                      ParagraphStyle("firma", fontSize=9, fontName="Helvetica",
                                     textColor=NEGRO, alignment=TA_CENTER)),
            Paragraph("Firma y sello<br/><br/><br/>___________________________<br/>Cliente / Receptor",
                      ParagraphStyle("firma", fontSize=9, fontName="Helvetica",
                                     textColor=NEGRO, alignment=TA_CENTER)),
        ]
    ]
    historia.append(
        Table(firma_data, colWidths=[ancho_col, ancho_col],
              style=TableStyle([("TOPPADDING", (0, 0), (-1, -1), 6),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
    )

    doc.build(historia)
    buffer.seek(0)
    return buffer.read()


# ─── EXPORTAR A DOCX ────────────────────────────────────────────────────────────

def _set_cell_bg(celda, hex_color: str):
    """Rellena el fondo de una celda de tabla Word."""
    tc_pr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def exportar_docx(contenido: str, tipo_doc: str = "documento", metadata: dict = None) -> bytes:
    """
    Genera un DOCX profesional a partir del texto generado por IA.

    Args:
        contenido: Texto del documento generado
        tipo_doc: "contrato" | "presupuesto" | "informe"
        metadata: dict con claves opcionales: empresa, cliente, fecha, numero

    Returns:
        bytes del DOCX listo para descargar
    """
    metadata = metadata or {}
    doc = Document()

    # ── Márgenes ────────────────────────────────────────────────────────────
    for seccion in doc.sections:
        seccion.top_margin = Cm(2.5)
        seccion.bottom_margin = Cm(2.5)
        seccion.left_margin = Cm(2.5)
        seccion.right_margin = Cm(2.5)

    # ── Cabecera azul ───────────────────────────────────────────────────────
    tabla_cab = doc.add_table(rows=1, cols=1)
    tabla_cab.style = "Table Grid"
    celda_cab = tabla_cab.cell(0, 0)
    _set_cell_bg(celda_cab, "1E3A5F")
    p_cab = celda_cab.paragraphs[0]
    p_cab.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cab = p_cab.add_run(tipo_doc.upper())
    run_cab.font.bold = True
    run_cab.font.size = Pt(16)
    run_cab.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p_cab.paragraph_format.space_before = Pt(8)
    p_cab.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # ── Metadatos ───────────────────────────────────────────────────────────
    campos_meta = [
        ("Empresa:", metadata.get("empresa", "")),
        ("Cliente:", metadata.get("cliente", "")),
        ("Fecha:", metadata.get("fecha", "")),
        ("Número:", metadata.get("numero", "")),
    ]
    campos_meta = [(k, v) for k, v in campos_meta if v]

    if campos_meta:
        tabla_meta = doc.add_table(rows=len(campos_meta), cols=2)
        tabla_meta.style = "Table Grid"
        for i, (clave, valor) in enumerate(campos_meta):
            celda_k = tabla_meta.cell(i, 0)
            celda_v = tabla_meta.cell(i, 1)
            _set_cell_bg(celda_k, "F4F6FA")
            _set_cell_bg(celda_v, "F4F6FA")
            p_k = celda_k.paragraphs[0]
            p_v = celda_v.paragraphs[0]
            run_k = p_k.add_run(clave)
            run_k.bold = True
            run_k.font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
            run_k.font.size = Pt(9)
            run_v = p_v.add_run(valor)
            run_v.font.size = Pt(9)
        doc.add_paragraph()

    # ── Línea separadora ────────────────────────────────────────────────────
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after = Pt(6)
    pPr = p_sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Cuerpo ──────────────────────────────────────────────────────────────
    bloques = _parsear_texto(contenido)

    for bloque in bloques:
        t = bloque["tipo"]
        if t == "espacio":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
        elif t == "h1":
            p = doc.add_heading(bloque["texto"], level=1)
            p.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            p.runs[0].font.size = Pt(13)
        elif t == "h2":
            p = doc.add_heading(bloque["texto"], level=2)
            p.runs[0].font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
            p.runs[0].font.size = Pt(11)
        elif t == "bullet":
            p = doc.add_paragraph(bloque["texto"], style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
        elif t == "numerado":
            p = doc.add_paragraph(bloque["texto"], style="List Number")
            p.paragraph_format.space_after = Pt(3)
        else:
            p = doc.add_paragraph(bloque["texto"])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(10)

    # ── Pie de firmas ────────────────────────────────────────────────────────
    doc.add_paragraph()
    p_sep2 = doc.add_paragraph()
    pPr2 = p_sep2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    bottom2 = OxmlElement("w:bottom")
    bottom2.set(qn("w:val"), "single")
    bottom2.set(qn("w:sz"), "6")
    bottom2.set(qn("w:space"), "1")
    bottom2.set(qn("w:color"), "D0D7E3")
    pBdr2.append(bottom2)
    pPr2.append(pBdr2)
    doc.add_paragraph()

    tabla_firma = doc.add_table(rows=1, cols=2)
    for i, texto_firma in enumerate(["Representante empresa", "Cliente / Receptor"]):
        celda = tabla_firma.cell(0, i)
        p_f = celda.paragraphs[0]
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f.add_run("Firma y sello\n\n\n___________________________\n" + texto_firma).font.size = Pt(9)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()